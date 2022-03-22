import docx
from docx.shared import Pt
import os
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches

doc = docx.Document("template.docx")

sections = doc.sections
for section in sections:
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.5)


# Font style for full name will change name later prob
font_styles = doc.styles
font_charstyle = font_styles.add_style('style1', WD_STYLE_TYPE.CHARACTER)
font_object = font_charstyle.font
font_object.size = Pt(16)
font_object.name = "Times New Roman"

# Font style for personal info, may change name later
font_styles = doc.styles
font_charstyle = font_styles.add_style('style2', WD_STYLE_TYPE.CHARACTER)
font_object = font_charstyle.font
font_object.size = Pt(14)
font_object.name = 'Times New Roman'

font_styles = doc.styles
font_charstyle = font_styles.add_style('style3', WD_STYLE_TYPE.CHARACTER)
font_object = font_charstyle.font
font_object.size = Pt(14)
font_object.name = 'Calibri (Body)'

font_styles = doc.styles
font_charstyle = font_styles.add_style('style4', WD_STYLE_TYPE.CHARACTER)
font_object = font_charstyle.font
font_object.size = Pt(12)
font_object.name = 'Calibri (Body)'

table = doc.add_table(rows=4, cols=1)
table.style = 'Plain Table 4'

def resume_creator():
    add_info()
    job()
    education()
    certifications()
    skills()
    answer = input("Enter file name to save as: ")
    doc.save(f"{answer.lower()}.docx")
    os.system(f"start {answer.lower()}.docx")

def job():
    print("Lets collect your job information.")
    company = input("Enter name of company: ")
    job_title = input("Enter job title: ")
    start_date = input("Enter start date (Month Year): ")
    end_date = input("Enter end date (Month Year): ")
    print("Enter three job responsibilities.")
    jr1 = input("Job responsibility 1: ")
    jr2 = input("Job responsibility 2: ")
    jr3 = input("Job responsibility 3: ")

    job_1 = {
        # Dictionary containing info on job
        # JR means Job responsibility
        "Company": company,
        "Job": job_title,
        "Start Date": start_date,
        "End Date": end_date,
        "JR1": jr1,
        "JR2": jr2,
        "JR3": jr3
    }

    hdr_cells = table.rows[1].cells
    paragraph = hdr_cells[0].paragraphs[0]
    paragraph.add_run("WORK EXPERIENCE" + "\n", style="style2").bold = True
    paragraph.add_run("\n" + job_1["Job"].upper() + " ", style="style3").bold = True
    paragraph.add_run("\u2219" + " " + job_1["Company"].upper(), style="style3")
    paragraph.add_run("\n" + job_1["Start Date"].upper() + " " + "-" + " " + job_1["End Date"].upper(), style="style4").italic = True
    paragraph.add_run("\n" + "\u2022" + "   " + job_1["JR1"], style="style4")
    paragraph.add_run("\n" + "\u2022" + "   " + job_1["JR2"], style="style4")
    paragraph.add_run("\n" + "\u2022" + "   " + job_1["JR3"], style="style4")

    global run
    run = True
    while run:
        answer = input("Would you like to enter more job experience, you can add a maximum of 3. (yes/no) ")
        if answer.lower() == "yes":
            job2()
            answer = input("Would you like to enter more job experience, you can add a maximum of 3. (yes/no) ")
            while run:
                if answer.lower() == "yes":
                    job3()
                    run = False
                elif answer.lower() == "no":
                    run = False
                else:
                    print("Please enter 'yes' or 'no'.")
        elif answer.lower() == "no":
            run = False
        else:
            print("Please enter 'yes' or 'no'.")

def job2():
    company = input("Enter name of company: ")
    job_title = input("Enter job title: ")
    start_date = input("Enter start date (Month Year): ")
    end_date = input("Enter end date (Month Year): ")
    print("Enter three job responsibilities..")
    jr1 = input("Job responsibility 1: ")
    jr2 = input("Job responsibility 2: ")
    jr3 = input("Job responsibility 3: ")

    job_2 = {
        # Dictionary containing info on job
        # JR means Job responsibility
        "Company": company,
        "Job": job_title,
        "Start Date": start_date,
        "End Date": end_date,
        "JR1": jr1,
        "JR2": jr2,
        "JR3": jr3
    }

    hdr_cells = table.rows[1].cells
    paragraph = hdr_cells[0].paragraphs[0]
    paragraph.add_run("\n")
    paragraph.add_run("\n" + job_2["Job"].upper() + " ", style="style3").bold = True
    paragraph.add_run("\u2219" + " " + job_2["Company"].upper(), style="style3")
    paragraph.add_run("\n" + job_2["Start Date"].upper() + " " + "-" + " " + job_2["End Date"].upper(),
                      style="style4").italic = True
    paragraph.add_run("\n" + "\u2022" + "   " + job_2["JR1"], style="style4")
    paragraph.add_run("\n" + "\u2022" + "   " + job_2["JR2"], style="style4")
    paragraph.add_run("\n" + "\u2022" + "   " + job_2["JR3"], style="style4")

def job3():
    company = input("Enter name of company: ")
    job_title = input("Enter job title: ")
    start_date = input("Enter start date (Month Year): ")
    end_date = input("Enter end date (Month Year): ")
    print("Enter three job responsibilities.")
    jr1 = input("Job responsibility 1: ")
    jr2 = input("Job responsibility 2: ")
    jr3 = input("Job responsibility 3: ")

    job_3 = {
        # Dictionary containing info on job
        # JR means Job responsibility
        "Company": company,
        "Job": job_title,
        "Start Date": start_date,
        "End Date": end_date,
        "JR1": jr1,
        "JR2": jr2,
        "JR3": jr3
    }

    hdr_cells = table.rows[1].cells
    paragraph = hdr_cells[0].paragraphs[0]
    paragraph.add_run("\n")
    paragraph.add_run("\n" + job_3["Job"].upper() + " ", style="style3").bold = True
    paragraph.add_run("\u2219" + " " + job_3["Company"].upper(), style="style3")
    paragraph.add_run("\n" + job_3["Start Date"].upper() + " " + "-" + " " + job_3["End Date"].upper(),
                      style="style4").italic = True
    paragraph.add_run("\n" + "\u2022" + "   " + job_3["JR1"], style="style4")
    paragraph.add_run("\n" + "\u2022" + "   " + job_3["JR2"], style="style4")
    paragraph.add_run("\n" + "\u2022" + "   " + job_3["JR3"], style="style4")

def add_info():
    print("Lets add personal information.")
    first_name = input("Enter first name: ")
    last_name = input("Enter your last name: ")
    email = input("Enter your email: ")
    phone = input("Enter your phone number (000-000-0000) format: ")
    city = input("Enter your city: ")
    state = input("Enter your state: ")

    user_information = {
        # Dictionary containing personal information
        "First Name": first_name,
        "Last Name": last_name,
        "Email": email,
        "Phone Number": phone,
        "City": city,
        "State": state
    }
    hdr_cells = table.rows[0].cells
    paragraph = hdr_cells[0].paragraphs[0]
    paragraph.add_run(user_information["First Name"].title() + " " + user_information["Last Name"].title(),
                      style="style1").bold = True
    paragraph.add_run("\n" + user_information["Email"] + " " + "\u2219" + " ", style="style2")
    paragraph.add_run(user_information["Phone Number"] + " " + "\u2219" + " ", style="style2")
    paragraph.add_run(user_information["City"].title() + " " + "\u2219" + " ", style="style2")
    paragraph.add_run(user_information["State"].upper(), style="style2")


def education():
    print("Let's add some education.")
    school = input("Enter name of school: ")
    degree = input("Enter degree: ")
    start_date = input("Enter start date (Month Year): ")
    end_date = input("Enter end date (Month Year): ")
    print("Enter any academic achievements, gpa, etc.")
    sd1 = input("School description 1: ")
    sd2 = input("School description 2: ")

    education_1 = {
        # Dictionary containing info on education
        # SD means school description
        "School": school,
        "Degree": degree,
        "Start Date": start_date,
        "End Date": end_date,
        "SD1": sd1,
        "SD2": sd2,
    }

    hdr_cells = table.rows[2].cells
    paragraph = hdr_cells[0].paragraphs[0]
    paragraph.add_run("EDUCATION" + "\n", style="style2").bold = True
    paragraph.add_run("\n" + education_1["Degree"].upper() + " ", style="style3").bold = True
    paragraph.add_run("\u2219" + " " + education_1["School"].upper(), style="style3")
    paragraph.add_run("\n" + education_1["Start Date"].upper() + " " + "-" + " " + education_1["End Date"].upper(),
                      style="style4").italic = True
    paragraph.add_run("\n" + "\u2022" + "   " + education_1["SD1"], style="style4")
    paragraph.add_run("\n" + "\u2022" + "   " + education_1["SD2"], style="style4")

    while True:
        answer = input("Would you like to add more education? We can add a max of two. (yes/no): ")
        if answer.lower() == 'yes':
            school = input("Enter name of school: ")
            degree = input("Enter degree: ")
            start_date = input("Enter start date (Month Year): ")
            end_date = input("Enter end date (Month Year): ")
            print("Enter any academic achievements, gpa, etc.")
            sd1 = input("School description 1: ")
            sd2 = input("School description 2: ")

            education_2 = {
                # Dictionary containing info on education
                # SD means school description
                "School": school,
                "Degree": degree,
                "Start Date": start_date,
                "End Date": end_date,
                "SD1": sd1,
                "SD2": sd2,
            }

            hdr_cells = table.rows[2].cells
            paragraph = hdr_cells[0].paragraphs[0]
            paragraph.add_run("\n")
            paragraph.add_run("\n" + education_2["Degree"].upper() + " ", style="style3").bold = True
            paragraph.add_run("\u2219" + " " + education_2["School"].upper(), style="style3")
            paragraph.add_run(
                "\n" + education_2["Start Date"].upper() + " " + "-" + " " + education_2["End Date"].upper(),
                style="style4").italic = True
            paragraph.add_run("\n" + "\u2022" + "   " + education_2["SD1"], style="style4")
            paragraph.add_run("\n" + "\u2022" + "   " + education_2["SD2"], style="style4")
            break
        elif answer.lower() == 'no':
            break
        else:
            print("Please enter 'yes' or 'no'.")

def certifications():
    certifications = []
    while True:
        print("Would you like to enter any certifications? We can add a maximum of three.")
        answer = input("Enter (yes/no): ")
        if answer.lower() == 'yes':
            print("If you have less than three certifications, leave blank and press enter.")
            while True:
                cert1 = input("Enter certification: ")
                if cert1 == "":
                    break
                else:
                    certifications.append(cert1)
                    cert2 = input("Enter certification: ")
                    if cert2 == "":
                        break
                    else:
                        certifications.append(cert2)
                        cert3 = input("Enter certification: ")
                        if cert3 == "":
                            break
                        else:
                            certifications.append(cert3)
                            break
            global x
            x = 4
            table.add_row()
            hdr_cells = table.rows[3].cells
            paragraph = hdr_cells[0].paragraphs[0]
            number_of_certs = len(certifications)
            paragraph.add_run("CERTIFICATIONS" + "\n", style="style2").bold = True
            if number_of_certs == 1:
                paragraph.add_run("\n" + "\u2022" + "   " + certifications[0], style="style4")
                break
            elif number_of_certs == 2:
                paragraph.add_run("\n" + "\u2022" + "   " + certifications[0], style="style4")
                paragraph.add_run("\n" + "\u2022" + "   " + certifications[1], style="style4")
                break
            elif number_of_certs == 3:
                paragraph.add_run("\n" + "\u2022" + "   " + certifications[0], style="style4")
                paragraph.add_run("\n" + "\u2022" + "   " + certifications[1], style="style4")
                paragraph.add_run("\n" + "\u2022" + "   " + certifications[2], style="style4")
                break
        elif answer.lower() == 'no':
            x = 3
            break
        else:
            print("Please enter 'yes' or 'no'.")

def skills():
    skill_list = []
    print("Lets add some skills.")
    print("Leave blank and press enter when finished")
    while True:
        skills = input("Enter a skill: ")
        if skills == "":
            break
        else:
            skill_list.append(skills)

    hdr_cells = table.rows[x].cells
    paragraph = hdr_cells[0].paragraphs[0]
    paragraph.add_run("SKILLS" + "\n", style="style2").bold = True
    paragraph.add_run("" + "\n")
    for skill in skill_list:
        paragraph.add_run(skill.title() + " " + "\u2219" + " ", style="style4")


resume_creator()
